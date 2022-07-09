from django.shortcuts import render, redirect
from django.core.mail import EmailMessage
from django.conf import settings
from django.http import HttpResponse, FileResponse
from docx import Document
from numpy import mean
from pathlib import Path
from datetime import datetime, timedelta
# import PIL
import matplotlib.pyplot as plt
from stats.models import Regions, RegionalStats, RegionalRestrictions

from io import BytesIO

graphFigsize = (7, 6)


# Create your views here.
def report_page(request):
    context = {}
    return render(request, 'reports/create_report_page.html', context)


def __linear_regression__(x: list, y: list) -> tuple:
    meanX = sum(x) / len(x)
    meanY = sum(y) / len(y)
    b1 = sum([(x[i] - meanX) * (y[i] - meanY) for i in range(len(x))]) / sum(
        [(x[i] - meanX) ** 2 for i in range(len(x))])
    b0 = meanY - b1 * meanX
    return b0, b1


def create_report(request):
    # region Чтение переменных из запроса
    regionName = request.GET.get('region')
    region = Regions.objects.get(region=regionName)
    startDate = datetime.strptime(request.GET.get('startDate'), '%Y-%m-%d').date()
    endDate = datetime.strptime(request.GET.get('endDate'), '%Y-%m-%d').date()
    startDateString = startDate.strftime("%d.%m.%Y")
    endDateString = endDate.strftime("%d.%m.%Y")
    measuresOn = request.GET.get('measureAnalysis') == 'on'
    emailOn = request.GET.get('emailSending') == 'on'
    email = request.GET.get('email')
    # endregion
    reportFileName = f'static/files/reports/{region} {startDate}-{endDate} measures-{measuresOn}.docx'

    if not Path(reportFileName).exists():
        region_stats = RegionalStats.objects.filter(date__gte=startDate - timedelta(days=1), date__lte=endDate,
                                                    region=region).order_by('date')
        # print(region_stats)
        # print(len(region_stats))

        report = Document()
        styles = report.styles
        heading = report.add_heading(f'Отчет: COVID-19 в регионе {regionName} в период с '
                                     f'{startDateString} по {endDateString}', 0)
        heading.style = styles['Title']
        report.add_heading('Распространение', 1)

        # region Таблица распространения
        report.add_paragraph('В данной таблице представлена статистика распространения коронавируса в отчетный период')
        stats_table = report.add_table(rows=1, cols=6)
        stats_hdr_cells = stats_table.rows[0].cells
        # print(stats_hdr_cells)
        # print('Мучу хедер')
        # plt.show()
        stats_hdr_cells[0].text = 'Дата'
        stats_hdr_cells[1].text = 'Новых случаев'
        stats_hdr_cells[2].text = 'Госпитализировано'
        stats_hdr_cells[3].text = 'Вылечились'
        stats_hdr_cells[4].text = 'Умерли'
        stats_hdr_cells[5].text = 'Вакцинировано полностью (всего)'
        '''if region.region != 'Россия':
            for i in range(1, len(region_stats))[::-1]:
                region_stats[i].died = region_stats[i].died - region_stats[i - 1].died'''
        # print([region.died for region in region_stats], i)
        for i in range(1, len(region_stats)):
            stats = region_stats[i]
            # print('Забиваю статистикой')
            row_cells = stats_table.add_row().cells
            row_cells[0].text = str(stats.date.strftime("%d.%m.%Y"))
            row_cells[1].text = str(stats.new_cases)
            row_cells[2].text = str(stats.hospitalised)
            row_cells[3].text = str(stats.recovered)
            row_cells[4].text = str(stats.died)
            row_cells[5].text = str(stats.vaccinated_fully_cumulative)
        # endregion

        region_stats = region_stats[1:]
        datesGraphList = [stat.date for stat in region_stats]
        newCasesGraphList = [stat.new_cases for stat in region_stats]
        hospitalisedGraphList = [stat.hospitalised for stat in region_stats]
        recoveredGraphList = [stat.recovered for stat in region_stats]
        diedGraphList = [stat.died for stat in region_stats]
        max_cases = max(region_stats, key=lambda item: item.new_cases)

        # region Описание распространения (7-14 дней)
        if 7 < len(region_stats) <= 14:
            report.add_heading("Описание распространения", 2)
            rising = True
            fromDate = region_stats[0].date
            if region_stats[1].new_cases - region_stats[0].new_cases >= 0:
                casesDescription = report.add_paragraph(
                    f"Начиная с {region_stats[0].date.strftime('%d.%m.%Y')} количество "
                    f"новых случаев росло до ")
            else:
                casesDescription = report.add_paragraph(
                    f"Начиная с {region_stats[0].date.strftime('%d.%m.%Y')} количество "
                    f"новых случаев падало до ")
                rising = False
            for i in range(1, len(region_stats)):
                if region_stats[i].new_cases - region_stats[i - 1].new_cases >= 0 and not rising:
                    casesDescription.add_run(f"{region_stats[i - 1].date.strftime('%d.%m.%Y')}. Затем росло до ")
                    rising = True
                elif region_stats[i].new_cases - region_stats[i - 1].new_cases < 0 and rising:
                    if region_stats[i - 1].new_cases == max_cases:
                        casesDescription.add_run(
                            f"{region_stats[i - 1].date.strftime('%d.%m.%Y')}. Пик заболеваемости - "
                            f"{region_stats[i - 1].new_cases} случаев - пришелся на ")
                    casesDescription.add_run(f"{region_stats[i - 1].date.strftime('%d.%m.%Y')}. Затем падало до ")
                    rising = False
            casesDescription.add_run(region_stats[-1].date.strftime('%d.%m.%Y'))
        # endregion

        # region Графики распространения (>7 дней)
        if len(region_stats) > 7:
            report.add_heading("Графики", 2)

            plt.figure(figsize=graphFigsize)
            plt.xticks(rotation=20)
            plt.plot(datesGraphList, newCasesGraphList)
            plt.title(f'Новые случаи заражения COVID-19\nв регионе {regionName}')
            plt.ylabel('Новые случаи')
            plt.savefig(f'static/files/graphs/Случаи {regionName} {startDate} {endDate}.png')
            report.add_picture(f'static/files/graphs/Случаи {regionName} {startDate} {endDate}.png')
            plt.close()

            plt.figure(figsize=graphFigsize)
            plt.xticks(rotation=20)
            plt.plot(datesGraphList, hospitalisedGraphList)
            plt.title(f'Госпитализации в регионе {regionName}')
            plt.ylabel('Госпитализировано')
            plt.savefig(f'static/files/graphs/Госпитализировано {regionName} {startDate} {endDate}.png')
            report.add_picture(f'static/files/graphs/Госпитализировано {regionName} {startDate} {endDate}.png')
            plt.close()

            plt.figure(figsize=graphFigsize)
            plt.xticks(rotation=20)
            plt.plot(datesGraphList, recoveredGraphList)
            plt.title(f'Выздоровления от COVID-19\nв регионе {regionName}')
            plt.ylabel('Выздоровело')
            plt.savefig(f'static/files/graphs/Выздоровело {regionName} {startDate} {endDate}.png')
            report.add_picture(f'static/files/graphs/Выздоровело {regionName} {startDate} {endDate}.png')
            plt.close()

            plt.figure(figsize=graphFigsize)
            plt.xticks(rotation=20)
            plt.plot(datesGraphList, diedGraphList)
            plt.title(f'Смертей в регионе {regionName}')
            plt.ylabel('Умерших')
            plt.savefig(f'static/files/graphs/Умерло {regionName} {startDate} {endDate}.png')
            report.add_picture(f'static/files/graphs/Умерло {regionName} {startDate} {endDate}.png')
            plt.close()
        # endregion

        # region Общая информация (>3 дней)
        if len(region_stats) > 3:
            report.add_heading("Общая информация за отчетный период", 2)
            report.add_heading("Всего", 3)
            report.add_paragraph(f"Новых случаев: {sum(newCasesGraphList)}")
            report.add_paragraph(f"Госпитализаций: {sum(hospitalisedGraphList)}")
            report.add_paragraph(f"Выздоровлений: {sum(recoveredGraphList)}")
            report.add_paragraph(f"Смертей: {sum(diedGraphList)}")
            report.add_paragraph(
                f"Вакцинировано в начале периода ({startDateString}): {region_stats[0].vaccinated_fully_cumulative}")
            report.add_paragraph(
                f"Вакцинировано в конце периода ({endDateString}): {region_stats[-1].vaccinated_fully_cumulative}")
            report.add_paragraph(f"Коллективный иммунитет в конце периода: {region_stats[-1].collective_immunity}%")

            report.add_heading("Средние значения", 3)
            report.add_paragraph(f"Новых случаев: {round(mean(newCasesGraphList), 3)}")
            report.add_paragraph(f"Госпитализаций: {round(mean(hospitalisedGraphList), 3)}")
            report.add_paragraph(f"Выздоровлений: {round(mean(recoveredGraphList), 3)}")
            report.add_paragraph(f"Смертей: {round(mean(diedGraphList), 3)}")
            report.add_paragraph(f"Наибольшее количество новых случаев было зафиксировано "
                                 f"{max_cases.date.strftime('%d.%m.%Y')} - {max_cases.new_cases} случая")
        # endregion

        if measuresOn:
            report.add_heading("Анализ мер", 1)
            report.add_heading("Общий обзор действовавших мер", 2)
            region_restrictions = RegionalRestrictions.objects.filter(date__gte=startDate, date__lte=endDate,
                                                                      region=region).order_by('date')
            all_restrictions = list(set([rest.restriction for rest in region_restrictions]))
            restrictions_dates = [0] * len(all_restrictions)
            restrictions_dates_strings = [0] * len(all_restrictions)
            # print(restrictions_dates)
            # print(all_restrictions)
            # print(f'Тип в листе - {type(all_restrictions[0])}')
            report.add_paragraph(f"На протяжении рассматриваемого периода в регионе действовали следующие ограничения:")
            for i in range(len(all_restrictions)):
                restriction = all_restrictions[i]
                # print(f'Ограничение {restriction}')
                restriction_dates = list(
                    RegionalRestrictions.objects.filter(region=region, restriction_id=restriction.id).order_by(
                        'date').values('date'))
                restriction_dates = [restriction_dates[0]['date'], restriction_dates[-1]['date']]
                restrictions_dates[i] = restriction_dates
                restriction_dates_strings = [restriction_dates[0].strftime("%d.%m.%Y"),
                                             restriction_dates[1].strftime("%d.%m.%Y")]
                restrictions_dates_strings[i] = restriction_dates_strings
                # print(restriction_dates)
                report.add_paragraph(
                    f'{restriction} (с {restriction_dates_strings[0]} до {restriction_dates_strings[1]})',
                    style='List Bullet')
            report.add_paragraph(f"Всего на выбранном промежутке действовало {len(all_restrictions)} различных мер.")

            report.add_heading("Анализ эффективности", 2)
            report.add_paragraph(f"В начале отчетного периода ({startDateString}) действовали следующие меры:")
            newMeasures = False
            gotMeasures = False
            for restr, dates in zip(all_restrictions, restrictions_dates):
                name = restr.description
                if dates[0] <= startDate:
                    report.add_paragraph(name, style='List Bullet')
                    gotMeasures = True
                else:
                    newMeasures = True
            if not gotMeasures:
                report.add_paragraph('Никаких ограничений.')
            if newMeasures:
                report.add_paragraph(f"Далее вводились ({startDateString}) следующие меры:")
                startDates = [item[0] for item in restrictions_dates]
                endDates = [item[1] for item in restrictions_dates]
                # print(f'lenallrestr - {len(all_restrictions)}')
                # print(f'lendates - {len(datesGraphList)}')
                for date in datesGraphList:
                    if date in startDates:
                        index = startDates.index(date)
                        restriction = all_restrictions[index].description.replace('/', '.')
                        # print(f'{date} - введено ограничение {restriction}')
                        report.add_paragraph(f'{date} - введено ограничение: {restriction}')
                        if (date - startDate).days >= 10 and (endDate - date).days >= 10:
                            index = datesGraphList.index(date)
                            analysisGraphDates = datesGraphList[index - 10:index + 10]
                            analysisXaxis = list(range(7))
                            # print(f'lengths - {len(analysisDates)}, {len(analysisXaxis)}')
                            analysisCases = newCasesGraphList[index - 7:index]
                            analysisGraphCases = newCasesGraphList[index - 10:index + 10]
                            # print(len(analysisDates), analysisDates)
                            # print(analysisXaxis)
                            # print(analysisCases)
                            b0, b1 = __linear_regression__(analysisXaxis, analysisCases)
                            regressionDates = datesGraphList[index-7:index+7]
                            regressionValues = [x*b1 + b0 for x in range(0, 14)]
                            # print()
                            # print(regressionDates)
                            # print(regressionValues)
                            # print()
                            plt.figure(figsize=graphFigsize)
                            plt.xticks(rotation=20)
                            plt.plot(analysisGraphDates, analysisGraphCases, label="Реальные значения")
                            plt.plot(regressionDates, regressionValues, 'r.--', label="Модель по неделе до введения")
                            plt.plot(datesGraphList[index], newCasesGraphList[index], 'go', label="Введение ограничения")
                            if len(restriction) > 50:
                                restriction = restriction[:47] + '...'
                            # print(f'type restr - {type(restriction)}')
                            plt.title(f'Введение меры\n{restriction}\nв регионе {regionName}')
                            plt.ylabel('Новые случаи заражения')
                            plt.legend()
                            plt.savefig(f'static/files/graphs/{regionName} {restriction} {date}.png')
                            report.add_picture(f'static/files/graphs/{regionName} {restriction} {date}.png')
                            plt.close()

        report.save(reportFileName)
    else:
        print('Данный отчет уже существует в системе')

    if emailOn:
        print('Отправка на почту')
        f = BytesIO()
        report.save(f)
        subject = f'Отчет по COVID-19 в регионе {regionName}'
        mail = EmailMessage(subject=subject, body='',
                            from_email=settings.EMAIL_HOST_USER, to=[email],
                            attachments=[["Отчет.docx", f.getvalue(),
                                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]])
        mail.send()
    response = FileResponse(open(reportFileName, 'rb'))
    return response
